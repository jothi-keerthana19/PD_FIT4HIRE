import os
import sys
import re
import argparse
import base64
from pathlib import Path
import requests
import PyPDF2
from docx import Document
import fitz  # PyMuPDF
from PIL import Image
import pytesseract
import io
import json
from groq import Groq

# Initialize Groq client
client = Groq(api_key="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")

# Job role templates with associated keywords and skills
JOB_ROLES = {
    "Software Developer": {
        "keywords": ["software development", "programming", "coding", "algorithms", "debugging"],
        "skills": ["Python", "Java", "JavaScript", "C++", "SQL", "Git", "REST API", "AWS", "Docker", "CI/CD"]
    },
    "Data Scientist": {
        "keywords": ["data analysis", "machine learning", "statistics", "data visualization", "research"],
        "skills": ["Python", "R", "SQL", "TensorFlow", "PyTorch", "Pandas", "NumPy", "Scikit-learn", "Tableau", "Big Data"]
    },
    "Project Manager": {
        "keywords": ["project management", "team leadership", "planning", "budgeting", "risk management"],
        "skills": ["Agile", "Scrum", "JIRA", "MS Project", "PMP", "stakeholder management", "resource allocation"]
    },
    "Marketing Specialist": {
        "keywords": ["marketing", "branding", "campaigns", "market research", "customer acquisition"],
        "skills": ["SEO", "SEM", "social media", "content marketing", "Google Analytics", "Adobe Creative Suite"]
    },
    "UX/UI Designer": {
        "keywords": ["user experience", "interface design", "wireframing", "prototyping", "usability testing"],
        "skills": ["Figma", "Sketch", "Adobe XD", "InVision", "HTML/CSS", "user research", "information architecture"]
    },
    "Frontend Developer": {
        "keywords": ["frontend development", "web development", "user interface", "responsive design", "browser compatibility"],
        "skills": ["React.js", "JavaScript", "HTML5", "CSS3", "TypeScript", "Redux", "Vue.js", "Angular", "Webpack", "SASS/LESS"]
    },
    "Backend Developer": {
        "keywords": ["backend development", "server-side", "api development", "database design", "system architecture"],
        "skills": ["Node.js", "Python", "Java", "SQL", "RESTful APIs", "Database Design", "MongoDB", "PostgreSQL", "Redis", "Microservices"]
    },
    "DevOps Engineer": {
        "keywords": ["devops", "automation", "infrastructure", "deployment", "monitoring"],
        "skills": ["Docker", "Kubernetes", "CI/CD", "AWS/Azure/GCP", "Terraform", "Jenkins", "Ansible", "Linux", "Shell Scripting", "Git"]
    },
    "Machine Learning Engineer": {
        "keywords": ["machine learning", "deep learning", "model development", "data modeling", "algorithm optimization"],
        "skills": ["Python", "TensorFlow", "PyTorch", "Scikit-learn", "Deep Learning", "NLP", "Computer Vision", "MLOps", "Statistical Analysis", "Feature Engineering"]
    },
    "Cloud Architect": {
        "keywords": ["cloud architecture", "system design", "infrastructure", "scalability", "cloud migration"],
        "skills": ["AWS", "Azure", "Google Cloud", "Kubernetes", "Microservices", "Security", "Terraform", "Docker", "CI/CD", "Serverless"]
    },
    "Security Engineer": {
        "keywords": ["security", "cybersecurity", "threat analysis", "vulnerability assessment", "security architecture"],
        "skills": ["Network Security", "Penetration Testing", "Security Tools", "Cryptography", "SIEM", "Incident Response", "Security Auditing", "Risk Assessment", "Compliance", "Cloud Security"]
    },
    "QA Engineer": {
        "keywords": ["quality assurance", "testing", "test automation", "bug tracking", "test planning"],
        "skills": ["Selenium", "API Testing", "Test Automation", "Manual Testing", "Performance Testing", "Test Cases", "Bug Tracking", "Regression Testing", "CI/CD", "Test Frameworks"]
    }
}

def extract_text_from_pdf(file_path):
    """Extract text from PDF files, including scanned PDFs."""
    text = ""
    images = []
    
    # Try using PyPDF2 first for text extraction
    try:
        with open(file_path, 'rb') as file:
            reader = PyPDF2.PdfReader(file)
            for page_num in range(len(reader.pages)):
                page = reader.pages[page_num]
                text += page.extract_text() + "\n"
    except Exception as e:
        print(f"PyPDF2 extraction error: {e}")
    
    # If text is limited or images are present, try PyMuPDF (fitz) for better extraction
    try:
        doc = fitz.open(file_path)
        if len(text.strip()) < 100:  # If PyPDF2 didn't get much text
            text = ""  # Reset text to use fitz extraction instead
            for page_num in range(len(doc)):
                page = doc[page_num]
                text += page.get_text() + "\n"
        
        # Extract images regardless of text content
        for page_num in range(len(doc)):
            page = doc[page_num]
            image_list = page.get_images(full=True)
            for img_index, img in enumerate(image_list):
                xref = img[0]
                base_image = doc.extract_image(xref)
                image_bytes = base_image["image"]
                
                # Convert image bytes to PIL Image
                image = Image.open(io.BytesIO(image_bytes))
                # Encode image to base64 for vision model
                buffered = io.BytesIO()
                image.save(buffered, format="PNG")
                img_base64 = base64.b64encode(buffered.getvalue()).decode('utf-8')
                
                # Use Groq vision model to analyze the image
                try:
                    vision_response = client.chat.completions.create(
                        model="llama-3.2-11b-vision-preview",
                        messages=[
                            {
                                "role": "user",
                                "content": [
                                    {
                                        "type": "text",
                                        "text": "Extract and summarize all text and relevant information from this resume image. Include any visible text, section headers, and structured content."
                                    },
                                    {
                                        "type": "image_url",
                                        "image_url": f"data:image/png;base64,{img_base64}"
                                    }
                                ]
                            }
                        ],
                        temperature=0.1,
                        max_tokens=1000
                    )
                    
                    # Add vision model's text extraction to the overall text
                    if vision_response.choices:
                        text += "\n" + vision_response.choices[0].message.content
                except Exception as e:
                    print(f"Vision model error: {e}")
                
                images.append(image)
        
        doc.close()
    except Exception as e:
        print(f"PyMuPDF extraction error: {e}")
    
    return text, images

def extract_text_from_docx(file_path):
    """Extract text from DOCX files."""
    doc = Document(file_path)
    text = ""
    for paragraph in doc.paragraphs:
        text += paragraph.text + "\n"
    
    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text += cell.text + "\n"
    
    return text, []  # No image extraction for docx for simplicity

def extract_text_from_image(image_path):
    """Extract text from image files using OCR."""
    try:
        image = Image.open(image_path)
        text = pytesseract.image_to_string(image)
        return text, [image]
    except Exception as e:
        print(f"Image OCR error: {e}")
        return "", []

def extract_text_from_images(images):
    """Extract text from a list of PIL images using OCR."""
    extracted_text = ""
    for img in images:
        try:
            text = pytesseract.image_to_string(img)
            extracted_text += text + "\n"
        except Exception as e:
            print(f"OCR processing error: {e}")
    
    return extracted_text

def calculate_ats_score(text, role, custom_skills=None):
    """Calculate ATS compatibility score based on job role and custom skills."""
    if role not in JOB_ROLES:
        return {
            "score": 0,
            "strengths": ["Invalid role selected"],
            "weaknesses": ["Please select a valid job role"]
        }
    
    # Get role keywords and skills
    role_data = JOB_ROLES[role]
    keywords = role_data["keywords"]
    skills = role_data["skills"] if not custom_skills else custom_skills
    
    text_lower = text.lower()
    
    # Enhanced keyword matching with frequency and context analysis
    keyword_scores = {}
    for keyword in keywords:
        keyword_lower = keyword.lower()
        # Count exact matches
        exact_matches = len(re.findall(r'\b' + re.escape(keyword_lower) + r'\b', text_lower))
        # Count partial matches (for compound words)
        partial_matches = len(re.findall(re.escape(keyword_lower), text_lower)) - exact_matches
        # Calculate weighted score
        keyword_scores[keyword] = min(exact_matches * 1.0 + partial_matches * 0.5, 3.0)  # Cap at 3
    
    keyword_score = (sum(keyword_scores.values()) / (len(keywords) * 3)) * 100
    
    # Enhanced skill matching with frequency and context analysis
    skill_scores = {}
    missing_skills = []
    for skill in skills:
        skill_lower = skill.lower()
        # Count exact matches
        exact_matches = len(re.findall(r'\b' + re.escape(skill_lower) + r'\b', text_lower))
        # Count partial matches
        partial_matches = len(re.findall(re.escape(skill_lower), text_lower)) - exact_matches
        # Calculate weighted score
        skill_score = min(exact_matches * 1.0 + partial_matches * 0.5, 2.0)  # Cap at 2
        
        if skill_score == 0:
            missing_skills.append(skill)
        else:
            skill_scores[skill] = skill_score
    
    skill_score = (sum(skill_scores.values()) / (len(skills) * 2)) * 100

    # Enhanced structure scoring with weighted sections
    structure_score = 0
    structure_strengths = []
    structure_weaknesses = []
    
    # Contact information (25%)
    contact_score = 0
    if re.search(r'[\w\.-]+@[\w\.-]+', text):  # Email
        contact_score += 10
        structure_strengths.append("Email present")
    else:
        structure_weaknesses.append("No email detected")
        
    if re.search(r'\b(?:\+\d{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}\b', text):  # Phone
        contact_score += 10
        structure_strengths.append("Phone number present")
    else:
        structure_weaknesses.append("No phone number detected")
    
    if re.search(r'linkedin\.com/\w+', text):  # LinkedIn
        contact_score += 5
        structure_strengths.append("LinkedIn profile present")
    else:
        structure_weaknesses.append("No LinkedIn profile detected")
    
    # Education section (20%)
    education_score = 0
    education_pattern = r'\b(education|university|college|degree|bachelor|master|phd)\b'
    if re.search(education_pattern, text_lower):
        matches = len(re.findall(education_pattern, text_lower))
        education_score = min(matches * 5, 20)
        structure_strengths.append("Education section present")
    else:
        structure_weaknesses.append("No education section detected")
    
    # Experience section (35%)
    experience_score = 0
    experience_pattern = r'\b(experience|work|job|position|role|employment)\b'
    if re.search(experience_pattern, text_lower):
        matches = len(re.findall(experience_pattern, text_lower))
        experience_score = min(matches * 7, 35)
        structure_strengths.append("Experience section present")
    else:
        structure_weaknesses.append("No experience section detected")
    
    # Additional sections (20%)
    additional_score = 0
    sections = {
        "skills": r'\b(skills|expertise|competencies)\b',
        "projects": r'\b(projects|portfolio)\b',
        "certifications": r'\b(certifications|certificates|accreditations)\b',
        "achievements": r'\b(achievements|accomplishments|awards)\b'
    }
    
    for section, pattern in sections.items():
        if re.search(pattern, text_lower):
            additional_score += 5
            structure_strengths.append(f"{section.title()} section present")
    
    structure_score = contact_score + education_score + experience_score + additional_score

    # Calculate weighted overall score
    overall_score = (keyword_score * 0.3) + (skill_score * 0.5) + (structure_score * 0.2)
    
    strengths = []
    weaknesses = []
    
    # Add keyword matches to strengths
    if keyword_score >= 70:
        strengths.append(f"Strong presence of {role} keywords ({sum(keyword_scores.values())}/{len(keywords)*3} matched)")
    elif keyword_score >= 40:
        strengths.append(f"Moderate presence of {role} keywords ({sum(keyword_scores.values())}/{len(keywords)*3} matched)")
    else:
        weaknesses.append(f"Poor keyword optimization for {role} ({sum(keyword_scores.values())}/{len(keywords)*3} matched)")
    
    # Add skill matches to strengths or weaknesses
    if skill_score >= 70:
        strengths.append(f"Strong presence of relevant skills ({sum(skill_scores.values())}/{len(skills)*2} matched)")
    elif skill_score >= 40:
        strengths.append(f"Moderate presence of relevant skills ({sum(skill_scores.values())}/{len(skills)*2} matched)")
    else:
        weaknesses.append(f"Insufficient skills for {role} ({sum(skill_scores.values())}/{len(skills)*2} matched)")
    
    # Add structure strengths and weaknesses
    strengths.extend(structure_strengths)
    weaknesses.extend(structure_weaknesses)
    
    # Add missing skills to weaknesses
    if missing_skills:
        missing_skills_str = ", ".join(missing_skills)
        weaknesses.append(f"Missing important skills: {missing_skills_str}")
    
    return {
        "score": round(overall_score, 1),
        "strengths": strengths,
        "weaknesses": weaknesses,
        "keyword_score": round(keyword_score, 1),
        "skill_score": round(skill_score, 1),
        "structure_score": structure_score,
        "missing_skills": missing_skills
    }

def analyze_resume(text, job_role, custom_skills=None):
    """Analyze the resume text using Groq API from a recruiter's perspective."""
    ats_analysis = calculate_ats_score(text, job_role, custom_skills)
    
    # Prepare role-specific context for the analysis
    role_context = ""
    if job_role in JOB_ROLES:
        skills_list = ", ".join(JOB_ROLES[job_role]["skills"])
        keywords_list = ", ".join(JOB_ROLES[job_role]["keywords"])
        role_context = f"""
For reference, required skills for a {job_role} position include: {skills_list}
Key indicators for this role include: {keywords_list}
        """
    
    if custom_skills:
        role_context += f"\nAdditional required skills for this position: {', '.join(custom_skills)}"

    prompt = f"""
As a technical recruiter, analyze this resume for a {job_role} position. Provide a detailed evaluation focusing on:

Candidate Overview:
- Name and Contact Information
- Years of relevant experience
- Current role and company
- Location and relocation status (if mentioned)

Technical Qualification Assessment:
- Core technical skills alignment with {job_role} requirements
- Technical skill proficiency levels (if indicated)
- Notable technical achievements or projects
- Missing critical skills for the role

Professional Experience:
- Career progression and growth
- Company profile/tier of previous employers
- Project scale and impact
- Team leadership or management experience
- Industry-specific experience

Education and Certifications:
- Relevant degrees and certifications
- Educational institution reputation
- Graduation dates and timeline
- Continuing education or professional development

Red Flags and Concerns:
- Employment gaps
- Job hopping patterns
- Mismatched skills or experience
- Unclear or vague descriptions
- Missing key information

Candidate Strengths:
- Standout achievements
- Unique qualifications
- Industry recognition or awards
- Notable projects or contributions

Interview Recommendations:
- Key areas to probe during technical interviews
- Specific technical skills to verify
- Experience areas requiring clarification
- Suggested technical assessment focus

Hiring Recommendation:
- Suitability for the {job_role} position
- Experience level match (junior/mid/senior)
- Potential alternative roles if not an exact match
- Suggested next steps in the hiring process

ATS Compatibility:
- ATS Score: {ats_analysis["score"]}%
- Key Matches: {", ".join(ats_analysis["strengths"])}
- Gaps: {", ".join(ats_analysis["weaknesses"])}
- Missing Required Skills: {", ".join(ats_analysis["missing_skills"])}

{role_context}

Resume Text:
{text}
"""

    try:
        response = client.chat.completions.create(
            model="llama3-70b-8192",
            messages=[
                {"role": "system", "content": f"You are a professional technical recruiter specializing in {job_role} positions. Provide thorough, objective analysis with hiring recommendations."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.1,
            max_tokens=4096
        )
        return response.choices[0].message.content, ats_analysis
    except Exception as e:
        print(f"Groq API error: {e}")
        return f"Error analyzing resume: {str(e)}", ats_analysis

def analyze_image_with_vision(image_path, job_role):
    """
    Analyze an image using Groq's vision model with job-role specific prompting for hiring purposes.
    
    Args:
        image_path (str): Path to the image file
        job_role (str): Target job role for analysis
    """
    # Encode image to base64
    with open(image_path, "rb") as image_file:
        base64_image = base64.b64encode(image_file.read()).decode('utf-8')
    
    # Create role-specific prompt
    role_context = ""
    if job_role in JOB_ROLES:
        skills_list = ", ".join(JOB_ROLES[job_role]["skills"])
        keywords_list = ", ".join(JOB_ROLES[job_role]["keywords"])
        role_context = f"For a {job_role} position. Required skills include: {skills_list}. Key indicators include: {keywords_list}."

    prompt = f"""As a technical recruiter, analyze this resume image for {role_context}
    Extract and evaluate:
    1. Candidate Information:
       - Personal details and contact information
       - Current role and company
       - Years of experience
       - Location details

    2. Technical Qualifications:
       - Core technical skills and proficiency levels
       - Technical achievements and projects
       - Relevant certifications

    3. Professional Background:
       - Career progression
       - Company profiles
       - Project impacts
       - Leadership experience

    4. Education:
       - Degrees and institutions
       - Graduation dates
       - Relevant coursework

    5. Additional Qualifications:
       - Industry recognition
       - Notable projects
       - Professional development

    Focus on elements that indicate candidate suitability for the {job_role} position."""

    try:
        chat_completion = client.chat.completions.create(
            messages=[
                {
                    "role": "user",
                    "content": [
                        {"type": "text", "text": prompt},
                        {
                            "type": "image_url",
                            "image_url": {
                                "url": f"data:image/jpeg;base64,{base64_image}",
                            },
                        },
                    ],
                }
            ],
            model="llama-3.2-11b-vision-preview",
            temperature=0.1,
            max_tokens=1000
        )
        return chat_completion.choices[0].message.content
    except Exception as e:
        print(f"Vision model error: {e}")
        return ""

def process_resume(file_path, job_role, custom_skills=None):
    """Process the resume file and return analysis."""
    file_extension = os.path.splitext(file_path)[1].lower()
    
    text = ""
    images = []
    
    # For image files, use vision model directly
    if file_extension in ['.jpg', '.jpeg', '.png', '.bmp', '.tiff', '.tif']:
        vision_text = analyze_image_with_vision(file_path, job_role)
        text = vision_text
    else:
        # Extract text based on file type
        if file_extension in ['.pdf']:
            text, images = extract_text_from_pdf(file_path)
        elif file_extension in ['.docx', '.doc']:
            text, images = extract_text_from_docx(file_path)
        
        # If text extraction yielded little text but we have images, use vision model
        if len(text.strip()) < 100 and images:
            for idx, image in enumerate(images):
                # Save image temporarily
                temp_path = f"temp_image_{idx}.png"
                image.save(temp_path)
                vision_text = analyze_image_with_vision(temp_path, job_role)
                text += "\n" + vision_text
                os.remove(temp_path)  # Clean up temporary file

    # If we still don't have much text, return an error
    if len(text.strip()) < 100:
        return "Could not extract sufficient text from the resume. Please check the file format or quality.", None
    
    # Analyze the extracted text
    analysis, ats_score = analyze_resume(text, job_role, custom_skills)
    
    return {
        "extracted_text": text,
        "analysis": analysis,
        "ats_score": ats_score
    }

def get_resume_path():
    """Ask the user to input a file path if not provided by command line."""
    print("Please enter the path to your resume file:")
    path = input("> ").strip()
    
    # Remove quotes if the user included them
    if (path.startswith('"') and path.endswith('"')) or (path.startswith("'") and path.endswith("'")):
        path = path[1:-1]
    
    return path

def select_job_role():
    """Let the user select a job role for analysis."""
    print("\nAvailable job roles for analysis:")
    for i, role in enumerate(JOB_ROLES.keys(), 1):
        print(f"{i}. {role}")
    
    print("\nSelect a job role (enter the number) or enter 'c' to provide a custom role:")
    choice = input("> ").strip()
    
    if choice.lower() == 'c':
        print("Enter your custom job role:")
        custom_role = input("> ").strip()
        return custom_role
    
    try:
        choice_idx = int(choice) - 1
        if 0 <= choice_idx < len(JOB_ROLES):
            return list(JOB_ROLES.keys())[choice_idx]
    except ValueError:
        pass
    
    print("Invalid selection. Defaulting to 'Software Developer'.")
    return "Software Developer"

def add_custom_skills():
    """Let the user add custom skills for analysis."""
    print("\nWould you like to add custom skills for the analysis? (y/n)")
    choice = input("> ").strip().lower()
    
    if choice == 'y':
        print("Enter skills separated by commas (e.g., Python, Project Management, Communication):")
        skills_input = input("> ").strip()
        return [skill.strip() for skill in skills_input.split(',') if skill.strip()]
    
    return None

def display_results(result):
    """Format and display the analysis results."""
    if isinstance(result, dict) and "ats_score" in result:
        ats_score = result["ats_score"]
        
        print("\n" + "="*80)
        print(f"{'HIRING ASSESSMENT':^80}")
        print("="*80)
        print(f"\nATS Score: {ats_score['score']}%")
        print(f"Technical Skills Match: {ats_score['skill_score']}%")
        print(f"Role Keyword Match: {ats_score['keyword_score']}%")
        print(f"Resume Structure Score: {ats_score['structure_score']}%")
        
        print("\n" + "-"*40)
        print("CANDIDATE STRENGTHS:")
        for strength in ats_score['strengths']:
            print(f"✓ {strength}")
        
        print("\n" + "-"*40)
        print("AREAS OF CONCERN:")
        for weakness in ats_score['weaknesses']:
            print(f"✗ {weakness}")
        
        if ats_score['missing_skills']:
            print("\n" + "-"*40)
            print("MISSING REQUIRED SKILLS:")
            for skill in ats_score['missing_skills']:
                print(f"• {skill}")
        
        print("\n" + "="*80)
        print(f"{'DETAILED RECRUITMENT ANALYSIS':^80}")
        print("="*80)
        print(f"\n{result['analysis']}")
    else:
        print(result)

def main():
    parser = argparse.ArgumentParser(description="Analyze resumes and provide detailed feedback")
    parser.add_argument("file_path", nargs="?", help="Path to the resume file (PDF, DOCX, or image)")
    parser.add_argument("--role", "-r", help="Job role for analysis")
    parser.add_argument("--skills", "-s", help="Custom skills for analysis (comma-separated)")
    parser.add_argument("--output", "-o", help="Output file path for the analysis (default: stdout)")
    
    args = parser.parse_args()
    
    # If file_path is not provided as an argument, prompt the user for it
    file_path = args.file_path
    if not file_path:
        file_path = get_resume_path()
    
    # Check if file exists
    if not os.path.exists(file_path):
        print(f"Error: File not found - {file_path}")
        retry = input("Would you like to enter a different path? (y/n): ").lower()
        if retry == 'y':
            file_path = get_resume_path()
            if not os.path.exists(file_path):
                print(f"Error: File not found - {file_path}")
                sys.exit(1)
        else:
            sys.exit(1)
    
    # Get job role
    job_role = args.role
    if not job_role:
        job_role = select_job_role()
    
    # Get custom skills
    custom_skills = None
    if args.skills:
        custom_skills = [skill.strip() for skill in args.skills.split(',')]
    else:
        custom_skills = add_custom_skills()
    
    print(f"\nProcessing resume: {file_path}")
    print(f"Job Role: {job_role}")
    if custom_skills:
        print(f"Custom Skills: {', '.join(custom_skills)}")
    
    print("\nAnalyzing resume... This may take a moment.")
    result = process_resume(file_path, job_role, custom_skills)
    
    # Output results
    if args.output:
        with open(args.output, 'w', encoding='utf-8') as f:
            if isinstance(result, dict):
                json.dump(result, f, indent=2)
            else:
                f.write(result)
        print(f"Analysis saved to {args.output}")
    else:
        display_results(result)

if __name__ == "__main__":
    main()